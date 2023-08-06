from docx import Document


def get_lines(file_path):
    doc = Document(file_path)
    paragraph_iter = iter(doc.paragraphs)

    try:
        p = next(paragraph_iter)

        # Iterate through each paragraph, looking for headings
        while True:
            if p.style.style_id.startswith('Heading'):
                # When one is found, record it, and get the associated text entry
                heading = p.text
                entry = ''

                p = next(paragraph_iter)

                try:
                    # Get all of the 'normal' style paragraphs for the entry, until
                    # the next heading
                    while not p.style.style_id.startswith('Heading'):
                        entry += p.text.lower() + '\n'
                        p = next(paragraph_iter)
                except StopIteration:
                    # If a stop iteration is found during this time, we still
                    # want to return the current entry, so we catch the StopIteration
                    # here, so that it is still returned below, before the outer
                    # try/except block
                    break
                finally:
                    # Return a tuple of the heading and the entry
                    # The [:-1] takes off the trailing '\n' on the entry
                    yield heading, entry[:-1]
    except StopIteration:
        pass


def add_header(file_path, header: str):
    doc = Document(file_path)

    try:
        doc.add_heading(header, level=1)
    except KeyError:
        pass
    doc.add_paragraph()

    doc.save(file_path)


def new_doc(file_path):
    Document().save(file_path)
